# -*- coding: utf-8 -*-
"""
Created on Wed Apr 27 10:51:47 2022

@author: yugojim
"""
def creat_word(DocumentPath,title,fileName):
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches
        #print(DocumentPath)
        document = Document()
        document.add_heading(title,0)
        f = open(DocumentPath + fileName,encoding='UTF-8')#開啟資料檔案檔案
        lines = f.readlines() 
        for i in range(len(lines)):
            #print(lines[i])
            if lines[i]=='#\n':
                document.add_page_break()
            else:
                ParagraphString = lines[i].replace("\n", "")
                if lines[i][0].isdigit():
                    document.add_paragraph(ParagraphString,'Heading 1')
                    #if ParagraphString ==  '3、組織架構摘述':
                    #    document.add_picture('組織架構.png', width=Inches(6))
                    #if ParagraphString ==  '4、安裝指南':
                    #    document.add_picture('webflow.png', width=Inches(6))
                elif lines[i][1].isdigit() and lines[i][0] == '(' :
                    document.add_paragraph(ParagraphString,'List Bullet')
                else:
                    document.add_paragraph(ParagraphString)
        f.close()    
        document.save(str(DocumentPath) + fileName + '.docx')
        return ('OK')
    except:
        return ('NG')
    
