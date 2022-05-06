from django.shortcuts import render,get_object_or_404
from django.http import HttpResponse,HttpResponseRedirect
from django.urls import reverse
import pandas as pd
import pathlib
from . import Function
import os
DocumentPath = str(pathlib.Path().absolute()) + "/static/doc/"
def sdlc1(request):
    try:
        questiontype = request.POST['questiontype']
        question = request.POST['question']
        user = request.POST['user']
        scope = request.POST['scope']
        priority = request.POST['priority']
        PM = request.POST['PM']
        RD = request.POST['RD']
        UI = request.POST['UI']
        QA = request.POST['QA']
        OS = request.POST['OS']
        CPU = request.POST['CPU']
        RAM = request.POST['RAM']
        HD = request.POST['HD']
        result = request.POST['result']
        fileName='sdlc1'
        path = DocumentPath+fileName
        f = open(path, 'w', encoding='UTF-8')
        lines = []
        lines.append('1、定義問題\n')
        lines.append('(1)問題類型\n')
        lines.append('\t' + questiontype + '\n')
        lines.append('(2)問題描述\n')
        lines.append('\t' + question + '\n')
        lines.append('(3)使用者\n')
        lines.append('\t' + user + '\n')
        lines.append('(4)截止日期\n')
        lines.append('\t' + scope + '\n')
        lines.append('(5)優先度\n')
        lines.append('\t' + priority + '\n')
        lines.append('2、人力成本\n')
        lines.append('\t專案管理 PM : ' + PM + '\n')
        lines.append('\t程式開發 RD : ' + RD + '\n')
        lines.append('\t介面設計 UI : ' + UI + '\n')
        lines.append('\t測試人員 QA : ' + QA + '\n')
        lines.append('3、硬體資源\n')
        lines.append('\t作業系統 OS : ' + OS + '\n')
        lines.append('\t處理器 CPU :' + CPU + '\n')
        lines.append('\t記憶體 RAM : ' + RAM + '\n')
        lines.append('\t硬碟 HD : ' + HD + '\n')
        lines.append('4、預估成果\n')
        lines.append('\t' + result + '\n')
        #print(lines)
        f.writelines(lines)
        f.close()
        title='分析與計畫(Analysis and Planning)'
        FuncResult = Function.creat_word(DocumentPath,title,fileName)
        '''
        from docx import Document
        #print(DocumentPath)
        document = Document()
        document.add_heading(title,0)
        f = open(DocumentPath + fileName,encoding="utf-8")        #開啟資料檔案檔案
        lines = f.readlines() 
        for i in range(len(lines)):
            #print(lines[i])
            if lines[i]=='#\n':
                document.add_page_break()
            else:
                ParagraphString = lines[i].replace("\n", "")
                if lines[i][0].isdigit():
                    document.add_paragraph(ParagraphString,'Heading 1')
                    #if ParagraphString ==  '3、組織架構摘述':
                    #    document.add_picture('組織架構.png', width=Inches(6))
                    #if ParagraphString ==  '4、安裝指南':
                    #    document.add_picture('webflow.png', width=Inches(6))
                elif lines[i][1].isdigit():
                    document.add_paragraph(ParagraphString,'List Bullet')
                else:
                    document.add_paragraph(ParagraphString)
        f.close()    
        document.save(str(DocumentPath) + fileName + '.docx')
        print('OK')
        '''
        #print(FuncResult)
        context = {
                'FuncResult' : FuncResult,
                }
        return render(request, 'sdlc1.html', context)
    except:
        print('except')
        return render(request, 'sdlc1.html')
    
def sdlc2(request):
    try:
        return render(request, 'sdlc2.html')
    except:
        return render(request, 'sdlc2.html')  

def sdlc3(request):
    try:
        return render(request, 'sdlc3.html')
    except:
        return render(request, 'sdlc3.html')    

def sdlc4(request):
    try:
        return render(request, 'sdlc4.html')
    except:
        return render(request, 'sdlc4.html')

def sdlc5(request):
    try:
        return render(request, 'sdlc5.html')
    except:
        return render(request, 'sdlc5.html')

def sdlc6(request):
    try:
        return render(request, 'sdlc6.html')
    except:
        return render(request, 'sdlc6.html')    
    
def index(request):
    try:
        return render(request, 'index.html')
    except:
        return render(request, 'index.html')